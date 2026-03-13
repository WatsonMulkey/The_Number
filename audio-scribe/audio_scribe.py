"""
Audio Scribe — Local Offline Transcription Tool
Uses faster-whisper for transcription, FFmpeg for format conversion.
Retro terminal GUI matching Resume Tailor aesthetic.
"""

import atexit
import gc
import json
import os
import queue
import re
import shutil
import struct
import subprocess
import sys
import tempfile
import threading
import time
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext

# ── Theme Constants ──────────────────────────────────────────────────────────
BG_COLOR = "#1a1a1a"
FG_COLOR = "#00ff00"
TEXT_BG = "#0d0d0d"
ACCENT_COLOR = "#00aaff"
PLACEHOLDER_COLOR = "#666666"
ERROR_COLOR = "#ff4444"
WARN_COLOR = "#ffaa00"
MONO_FONT = ("Courier New", 10)
MONO_FONT_BOLD = ("Courier New", 10, "bold")
TITLE_FONT = ("Courier New", 7)
HEADER_FONT = ("Courier New", 12, "bold")

# ── Audio Format Constants ───────────────────────────────────────────────────
SUPPORTED_FORMATS = (
    ("Audio Files", "*.wav *.mp3 *.mp4 *.m4a *.caf *.ogg *.flac *.wma *.aac *.webm"),
    ("WAV", "*.wav"),
    ("MP3", "*.mp3"),
    ("M4A/AAC", "*.m4a *.aac"),
    ("All Files", "*.*"),
)

# ── Model Info ───────────────────────────────────────────────────────────────
MODELS = {
    "tiny": {"label": "tiny (75MB)", "size": "75MB"},
    "base": {"label": "base (150MB)", "size": "150MB"},
    "small": {"label": "small (500MB)", "size": "500MB"},
    "medium": {"label": "medium (1.5GB)", "size": "1.5GB"},
    "large-v3": {"label": "large (3GB)", "size": "3GB"},
}

# ── Preferences ──────────────────────────────────────────────────────────────
PREFS_DIR = Path.home() / ".audio_scribe"
PREFS_FILE = PREFS_DIR / "preferences.json"

ASCII_HEADER = r"""
  ╔═══════════════════════════════════════════════════════╗
  ║   █████╗ ██╗   ██╗██████╗ ██╗ ██████╗               ║
  ║  ██╔══██╗██║   ██║██╔══██╗██║██╔═══██╗              ║
  ║  ███████║██║   ██║██║  ██║██║██║   ██║              ║
  ║  ██╔══██║██║   ██║██║  ██║██║██║   ██║              ║
  ║  ██║  ██║╚██████╔╝██████╔╝██║╚██████╔╝              ║
  ║  ╚═╝  ╚═╝ ╚═════╝ ╚═════╝ ╚═╝ ╚═════╝              ║
  ║  ███████╗ ██████╗██████╗ ██╗██████╗ ███████╗         ║
  ║  ██╔════╝██╔════╝██╔══██╗██║██╔══██╗██╔════╝         ║
  ║  ███████╗██║     ██████╔╝██║██████╔╝█████╗           ║
  ║  ╚════██║██║     ██╔══██╗██║██╔══██╗██╔══╝           ║
  ║  ███████║╚██████╗██║  ██║██║██████╔╝███████╗         ║
  ║  ╚══════╝ ╚═════╝╚═╝  ╚═╝╚═╝╚═════╝ ╚══════╝         ║
  ║                                                       ║
  ║  Local Offline Transcription      [faster-whisper]    ║
  ╚═══════════════════════════════════════════════════════╝
"""


# ── Utility Functions ────────────────────────────────────────────────────────

def find_ffmpeg():
    """Find FFmpeg binary. Checks PATH, then common install locations."""
    path = shutil.which("ffmpeg")
    if path:
        return path

    # Common Windows install locations
    candidates = [
        Path.home() / "AppData/Local/Microsoft/WinGet/Packages",
        Path("C:/ffmpeg/bin"),
        Path("C:/Program Files/ffmpeg/bin"),
        Path("C:/tools/ffmpeg/bin"),
    ]

    # Search winget packages directory
    winget_dir = Path.home() / "AppData/Local/Microsoft/WinGet/Packages"
    if winget_dir.is_dir():
        for pkg in winget_dir.iterdir():
            if "ffmpeg" in pkg.name.lower():
                for ffmpeg_exe in pkg.rglob("ffmpeg.exe"):
                    return str(ffmpeg_exe)

    for d in candidates:
        exe = d / "ffmpeg.exe"
        if exe.is_file():
            return str(exe)

    return None


def format_timestamp(seconds):
    """Format seconds to HH:MM:SS,mmm (SRT format)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def format_timestamp_vtt(seconds):
    """Format seconds to HH:MM:SS.mmm (VTT format)."""
    return format_timestamp(seconds).replace(",", ".")


def format_duration(seconds):
    """Format seconds to human-readable duration."""
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    if minutes < 60:
        return f"{minutes}m {secs}s"
    hours = int(minutes // 60)
    mins = minutes % 60
    return f"{hours}h {mins}m {secs}s"


def sanitize_filename(name, max_len=100):
    """Sanitize a string for use as a filename."""
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = name.strip('. ')
    return name[:max_len] if name else "transcription"


def get_file_info(filepath, ffmpeg_path):
    """Get audio file info using FFmpeg probe."""
    try:
        ffprobe = ffmpeg_path.replace("ffmpeg", "ffprobe")
        if not os.path.isfile(ffprobe) and not shutil.which("ffprobe"):
            ffprobe = ffmpeg_path  # Fall back to ffmpeg -i for info
            result = subprocess.run(
                [ffprobe, "-i", filepath],
                capture_output=True, text=True, timeout=10,
            )
            output = result.stderr  # ffmpeg -i prints to stderr
            duration_match = re.search(r"Duration:\s*(\d+):(\d+):(\d+)\.(\d+)", output)
            if duration_match:
                h, m, s, cs = duration_match.groups()
                duration = int(h) * 3600 + int(m) * 60 + int(s) + int(cs) / 100
            else:
                duration = 0
            return {"duration": duration, "raw": output}

        result = subprocess.run(
            [ffprobe, "-v", "quiet", "-print_format", "json", "-show_format", filepath],
            capture_output=True, text=True, timeout=10,
        )
        info = json.loads(result.stdout)
        duration = float(info.get("format", {}).get("duration", 0))
        return {"duration": duration, "format": info.get("format", {}).get("format_long_name", "Unknown")}
    except Exception:
        return {"duration": 0, "format": "Unknown"}


def load_preferences():
    """Load user preferences from disk."""
    try:
        if PREFS_FILE.is_file():
            return json.loads(PREFS_FILE.read_text())
    except Exception:
        pass
    return {}


def save_preferences(prefs):
    """Save user preferences to disk."""
    try:
        PREFS_DIR.mkdir(parents=True, exist_ok=True)
        PREFS_FILE.write_text(json.dumps(prefs, indent=2))
    except Exception:
        pass


# ── Export Functions ──────────────────────────────────────────────────────────

def export_txt(segments, filepath):
    """Export as plain text with timestamps."""
    with open(filepath, "w", encoding="utf-8") as f:
        for seg in segments:
            ts = format_timestamp_vtt(seg["start"])
            f.write(f"[{ts}] {seg['text'].strip()}\n")


def export_srt(segments, filepath):
    """Export as SRT subtitle format."""
    with open(filepath, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, 1):
            f.write(f"{i}\n")
            f.write(f"{format_timestamp(seg['start'])} --> {format_timestamp(seg['end'])}\n")
            f.write(f"{seg['text'].strip()}\n\n")


def export_vtt(segments, filepath):
    """Export as WebVTT format."""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("WEBVTT\n\n")
        for i, seg in enumerate(segments, 1):
            f.write(f"{i}\n")
            f.write(f"{format_timestamp_vtt(seg['start'])} --> {format_timestamp_vtt(seg['end'])}\n")
            f.write(f"{seg['text'].strip()}\n\n")


def export_json(segments, filepath, metadata=None):
    """Export as structured JSON."""
    data = {
        "metadata": metadata or {},
        "segments": [
            {
                "index": i,
                "start": seg["start"],
                "end": seg["end"],
                "text": seg["text"].strip(),
            }
            for i, seg in enumerate(segments, 1)
        ],
        "full_text": " ".join(seg["text"].strip() for seg in segments),
    }
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def export_docx(segments, filepath, metadata=None):
    """Export as formatted DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading("Audio Scribe — Transcription", level=1)
    title.runs[0].font.color.rgb = RGBColor(0, 170, 255)

    # Metadata
    if metadata:
        meta_para = doc.add_paragraph()
        for key, val in metadata.items():
            run = meta_para.add_run(f"{key}: {val}\n")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    # Segments
    for seg in segments:
        ts = format_timestamp_vtt(seg["start"])
        para = doc.add_paragraph()
        ts_run = para.add_run(f"[{ts}]  ")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0, 170, 255)
        text_run = para.add_run(seg["text"].strip())
        text_run.font.size = Pt(11)

    doc.save(filepath)


# ── Transcription Worker ─────────────────────────────────────────────────────

def transcription_worker(
    audio_path, model_name, language, ffmpeg_path,
    msg_queue, cancel_event, device="cpu",
):
    """Run transcription in a background thread. Communicates via queue."""
    temp_wav = None
    try:
        # Step 1: Convert to 16kHz mono WAV if needed
        ext = Path(audio_path).suffix.lower()
        if ext != ".wav":
            msg_queue.put(("status", "Converting audio format..."))
            temp_wav = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
            result = subprocess.run(
                [ffmpeg_path, "-i", audio_path, "-ar", "16000", "-ac", "1",
                 "-y", temp_wav],
                capture_output=True, text=True, timeout=300,
            )
            if result.returncode != 0:
                raise RuntimeError(f"FFmpeg conversion failed: {result.stderr[:500]}")
            work_path = temp_wav
        else:
            work_path = audio_path

        if cancel_event.is_set():
            msg_queue.put(("cancelled", None))
            return

        # Step 2: Load model
        msg_queue.put(("status", f"Loading {model_name} model (first run downloads it)..."))

        try:
            from faster_whisper import WhisperModel
            model = WhisperModel(model_name, device=device, compute_type="int8")
        except RuntimeError as e:
            if "out of memory" in str(e).lower() or "CUDA" in str(e):
                msg_queue.put(("status", "GPU OOM — falling back to CPU..."))
                model = WhisperModel(model_name, device="cpu", compute_type="int8")
            else:
                raise

        if cancel_event.is_set():
            msg_queue.put(("cancelled", None))
            return

        # Step 3: Transcribe
        lang = None if language == "auto" else language
        msg_queue.put(("status", "Transcribing..."))

        segments_iter, info = model.transcribe(
            work_path, beam_size=5, language=lang,
            vad_filter=True, vad_parameters=dict(min_silence_duration_ms=500),
        )

        # Estimate total segments from duration
        file_info = get_file_info(audio_path, ffmpeg_path)
        est_total = max(1, int(file_info["duration"] / 5))  # ~5s per segment

        segments = []
        for i, seg in enumerate(segments_iter):
            if cancel_event.is_set():
                msg_queue.put(("cancelled", None))
                return
            segment_data = {
                "start": seg.start,
                "end": seg.end,
                "text": seg.text,
            }
            segments.append(segment_data)
            msg_queue.put(("segment", {
                "index": i + 1,
                "estimated_total": est_total,
                "segment": segment_data,
            }))

        if not segments:
            msg_queue.put(("no_speech", None))
            return

        # Build metadata
        metadata = {
            "source_file": os.path.basename(audio_path),
            "duration": format_duration(file_info["duration"]),
            "model": model_name,
            "language_detected": info.language,
            "language_probability": f"{info.language_probability:.1%}",
            "word_count": sum(len(s["text"].split()) for s in segments),
            "segment_count": len(segments),
            "transcribed_at": datetime.now().isoformat(),
        }

        msg_queue.put(("complete", {"segments": segments, "metadata": metadata}))

    except Exception as e:
        msg_queue.put(("error", str(e)))
    finally:
        # Clean up temp file
        if temp_wav and os.path.isfile(temp_wav):
            try:
                gc.collect()
                time.sleep(0.2)
                os.unlink(temp_wav)
            except PermissionError:
                pass  # Windows file locking — will be cleaned by OS


# ── Main Application ─────────────────────────────────────────────────────────

class AudioScribeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Audio Scribe")
        self.root.configure(bg=BG_COLOR)
        self.root.minsize(800, 700)

        # State
        self.ffmpeg_path = find_ffmpeg()
        self.audio_path = None
        self.segments = []
        self.metadata = {}
        self.msg_queue = queue.Queue()
        self.cancel_event = threading.Event()
        self.worker_thread = None
        self.temp_files = []
        self.prefs = load_preferences()
        self.start_time = None

        # Register cleanup
        atexit.register(self._cleanup_temp_files)
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

        # Check FFmpeg on startup
        if not self.ffmpeg_path:
            self.root.after(100, self._show_ffmpeg_error)

        self._build_gui()
        self._poll_queue()

    def _show_ffmpeg_error(self):
        messagebox.showerror(
            "FFmpeg Not Found",
            "Audio Scribe requires FFmpeg for audio format conversion.\n\n"
            "Install it with:\n"
            "  winget install Gyan.FFmpeg\n\n"
            "Or download from https://ffmpeg.org\n\n"
            "After installing, restart Audio Scribe.",
        )

    def _build_gui(self):
        # ── ASCII Header ──
        header_frame = tk.Frame(self.root, bg=BG_COLOR)
        header_frame.pack(fill=tk.X, padx=10)
        header_label = tk.Label(
            header_frame, text=ASCII_HEADER, font=TITLE_FONT,
            bg=BG_COLOR, fg=FG_COLOR, justify=tk.LEFT,
        )
        header_label.pack(anchor=tk.W)

        # ── Controls Frame ──
        controls = tk.Frame(self.root, bg=BG_COLOR)
        controls.pack(fill=tk.X, padx=15, pady=(0, 5))

        # File selection row
        file_row = tk.Frame(controls, bg=BG_COLOR)
        file_row.pack(fill=tk.X, pady=3)

        tk.Label(
            file_row, text="FILE:", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ACCENT_COLOR,
        ).pack(side=tk.LEFT)

        self.file_label = tk.Label(
            file_row, text="  No file selected", font=MONO_FONT,
            bg=BG_COLOR, fg=PLACEHOLDER_COLOR, anchor=tk.W,
        )
        self.file_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))

        self.browse_btn = tk.Button(
            file_row, text="[ BROWSE ]", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=FG_COLOR, activebackground=BG_COLOR,
            activeforeground=ACCENT_COLOR, relief=tk.FLAT, cursor="hand2",
            command=self._browse_file,
        )
        self.browse_btn.pack(side=tk.RIGHT)

        # File info display
        self.file_info_label = tk.Label(
            controls, text="", font=MONO_FONT,
            bg=BG_COLOR, fg=PLACEHOLDER_COLOR, anchor=tk.W,
        )
        self.file_info_label.pack(fill=tk.X, pady=(0, 3))

        # Separator
        tk.Frame(controls, bg=FG_COLOR, height=1).pack(fill=tk.X, pady=5)

        # Model selection row
        model_row = tk.Frame(controls, bg=BG_COLOR)
        model_row.pack(fill=tk.X, pady=3)

        tk.Label(
            model_row, text="MODEL:", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ACCENT_COLOR,
        ).pack(side=tk.LEFT)

        self.model_var = tk.StringVar(value=self.prefs.get("model", "tiny"))
        for model_key, model_info in MODELS.items():
            rb = tk.Radiobutton(
                model_row, text=model_info["label"], variable=self.model_var,
                value=model_key, font=MONO_FONT, bg=BG_COLOR, fg=FG_COLOR,
                selectcolor=TEXT_BG, activebackground=BG_COLOR,
                activeforeground=ACCENT_COLOR, relief=tk.FLAT,
            )
            rb.pack(side=tk.LEFT, padx=(10, 0))

        # Language selection row
        lang_row = tk.Frame(controls, bg=BG_COLOR)
        lang_row.pack(fill=tk.X, pady=3)

        tk.Label(
            lang_row, text="LANG: ", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ACCENT_COLOR,
        ).pack(side=tk.LEFT)

        self.lang_var = tk.StringVar(value=self.prefs.get("language", "auto"))
        for lang_key, lang_label in [("auto", "Auto-detect"), ("en", "English")]:
            rb = tk.Radiobutton(
                lang_row, text=lang_label, variable=self.lang_var,
                value=lang_key, font=MONO_FONT, bg=BG_COLOR, fg=FG_COLOR,
                selectcolor=TEXT_BG, activebackground=BG_COLOR,
                activeforeground=ACCENT_COLOR, relief=tk.FLAT,
            )
            rb.pack(side=tk.LEFT, padx=(10, 0))

        # Separator
        tk.Frame(controls, bg=FG_COLOR, height=1).pack(fill=tk.X, pady=5)

        # Action buttons row
        btn_row = tk.Frame(controls, bg=BG_COLOR)
        btn_row.pack(fill=tk.X, pady=3)

        self.transcribe_btn = tk.Button(
            btn_row, text="[ TRANSCRIBE ]", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=FG_COLOR, activebackground=BG_COLOR,
            activeforeground=ACCENT_COLOR, relief=tk.FLAT, cursor="hand2",
            command=self._start_transcription,
        )
        self.transcribe_btn.pack(side=tk.LEFT)

        self.cancel_btn = tk.Button(
            btn_row, text="[ CANCEL ]", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ERROR_COLOR, activebackground=BG_COLOR,
            activeforeground=WARN_COLOR, relief=tk.FLAT, cursor="hand2",
            command=self._cancel_transcription, state=tk.DISABLED,
        )
        self.cancel_btn.pack(side=tk.LEFT, padx=(15, 0))

        # Status label
        self.status_label = tk.Label(
            controls, text="Ready.", font=MONO_FONT,
            bg=BG_COLOR, fg=PLACEHOLDER_COLOR, anchor=tk.W,
        )
        self.status_label.pack(fill=tk.X, pady=(3, 0))

        # ── Output Area ──
        output_frame = tk.Frame(self.root, bg=BG_COLOR)
        output_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(5, 5))

        tk.Label(
            output_frame, text="╔══ OUTPUT ══╗", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ACCENT_COLOR, anchor=tk.W,
        ).pack(fill=tk.X)

        self.output_text = scrolledtext.ScrolledText(
            output_frame, font=MONO_FONT, bg=TEXT_BG, fg=FG_COLOR,
            insertbackground=FG_COLOR, relief=tk.FLAT, borderwidth=2,
            wrap=tk.WORD, state=tk.DISABLED,
        )
        self.output_text.pack(fill=tk.BOTH, expand=True)

        # ── Export Buttons ──
        export_frame = tk.Frame(self.root, bg=BG_COLOR)
        export_frame.pack(fill=tk.X, padx=15, pady=(0, 10))

        tk.Label(
            export_frame, text="EXPORT:", font=MONO_FONT_BOLD,
            bg=BG_COLOR, fg=ACCENT_COLOR,
        ).pack(side=tk.LEFT)

        export_buttons = [
            ("COPY", self._export_copy),
            ("TXT", self._export_txt),
            ("SRT", self._export_srt),
            ("VTT", self._export_vtt),
            ("JSON", self._export_json),
            ("DOCX", self._export_docx),
        ]

        self.export_btns = []
        for label, cmd in export_buttons:
            btn = tk.Button(
                export_frame, text=f"[ {label} ]", font=MONO_FONT_BOLD,
                bg=BG_COLOR, fg=FG_COLOR, activebackground=BG_COLOR,
                activeforeground=ACCENT_COLOR, relief=tk.FLAT, cursor="hand2",
                command=cmd, state=tk.DISABLED,
            )
            btn.pack(side=tk.LEFT, padx=(10, 0))
            self.export_btns.append(btn)

        # Stats label at bottom
        self.stats_label = tk.Label(
            self.root, text="", font=MONO_FONT,
            bg=BG_COLOR, fg=PLACEHOLDER_COLOR, anchor=tk.W,
        )
        self.stats_label.pack(fill=tk.X, padx=15, pady=(0, 5))

    # ── File Selection ───────────────────────────────────────────────────

    def _browse_file(self):
        initial_dir = self.prefs.get("last_dir", str(Path.home()))
        filepath = filedialog.askopenfilename(
            title="Select Audio File",
            initialdir=initial_dir,
            filetypes=SUPPORTED_FORMATS,
        )
        if not filepath:
            return

        self.audio_path = filepath
        self.prefs["last_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)

        filename = os.path.basename(filepath)
        size_mb = os.path.getsize(filepath) / (1024 * 1024)
        self.file_label.config(text=f"  {filename}", fg=FG_COLOR)

        # Get file info
        info_parts = [f"Size: {size_mb:.1f}MB"]
        if self.ffmpeg_path:
            finfo = get_file_info(filepath, self.ffmpeg_path)
            if finfo["duration"] > 0:
                info_parts.append(f"Duration: {format_duration(finfo['duration'])}")
            if "format" in finfo:
                info_parts.append(f"Format: {finfo['format']}")

        self.file_info_label.config(text="  " + "  |  ".join(info_parts))
        self.status_label.config(text="File loaded. Ready to transcribe.", fg=FG_COLOR)

    # ── Transcription Control ────────────────────────────────────────────

    def _start_transcription(self):
        if not self.audio_path:
            messagebox.showwarning("No File", "Please select an audio file first.")
            return

        if not self.ffmpeg_path:
            self._show_ffmpeg_error()
            return

        if not os.path.isfile(self.audio_path):
            messagebox.showerror("File Not Found", f"Cannot find:\n{self.audio_path}")
            return

        # Reset state
        self.segments = []
        self.metadata = {}
        self.cancel_event.clear()
        self._set_output("")
        self._set_transcribing(True)
        self.start_time = time.time()

        # Save model/language preference
        self.prefs["model"] = self.model_var.get()
        self.prefs["language"] = self.lang_var.get()
        save_preferences(self.prefs)

        # Launch worker
        self.worker_thread = threading.Thread(
            target=transcription_worker,
            args=(
                self.audio_path,
                self.model_var.get(),
                self.lang_var.get(),
                self.ffmpeg_path,
                self.msg_queue,
                self.cancel_event,
            ),
            daemon=True,
        )
        self.worker_thread.start()

    def _cancel_transcription(self):
        self.cancel_event.set()
        self.status_label.config(text="Cancelling...", fg=WARN_COLOR)

    def _set_transcribing(self, active):
        """Toggle button states for transcription active/inactive."""
        if active:
            self.transcribe_btn.config(state=tk.DISABLED)
            self.cancel_btn.config(state=tk.NORMAL)
            self.browse_btn.config(state=tk.DISABLED)
            for btn in self.export_btns:
                btn.config(state=tk.DISABLED)
        else:
            self.transcribe_btn.config(state=tk.NORMAL)
            self.cancel_btn.config(state=tk.DISABLED)
            self.browse_btn.config(state=tk.NORMAL)
            if self.segments:
                for btn in self.export_btns:
                    btn.config(state=tk.NORMAL)

    # ── Queue Polling ────────────────────────────────────────────────────

    def _poll_queue(self):
        """Poll message queue from worker thread. Runs on main thread via after()."""
        try:
            while True:
                msg_type, data = self.msg_queue.get_nowait()

                if msg_type == "status":
                    self.status_label.config(text=data, fg=ACCENT_COLOR)

                elif msg_type == "segment":
                    idx = data["index"]
                    est = data["estimated_total"]
                    seg = data["segment"]
                    self.segments.append(seg)

                    # Update progress
                    pct = min(99, int(idx / est * 100)) if est > 0 else 0
                    elapsed = time.time() - self.start_time if self.start_time else 0
                    self.status_label.config(
                        text=f"Segment {idx}/~{est}  ({pct}%)  [{format_duration(elapsed)}]",
                        fg=ACCENT_COLOR,
                    )

                    # Append to output
                    ts = format_timestamp_vtt(seg["start"])
                    self._append_output(f"[{ts}] {seg['text'].strip()}\n")

                elif msg_type == "complete":
                    self.segments = data["segments"]
                    self.metadata = data["metadata"]
                    elapsed = time.time() - self.start_time if self.start_time else 0
                    self._set_transcribing(False)
                    self.status_label.config(
                        text=f"Complete! {len(self.segments)} segments in {format_duration(elapsed)}",
                        fg=FG_COLOR,
                    )
                    # Stats
                    m = self.metadata
                    self.stats_label.config(
                        text=(
                            f"Words: {m.get('word_count', 0)}  |  "
                            f"Language: {m.get('language_detected', '?')} "
                            f"({m.get('language_probability', '?')})  |  "
                            f"Model: {m.get('model', '?')}  |  "
                            f"Elapsed: {format_duration(elapsed)}"
                        ),
                    )

                elif msg_type == "no_speech":
                    self._set_transcribing(False)
                    self.status_label.config(
                        text="No speech detected in audio file.", fg=WARN_COLOR,
                    )
                    self._set_output("[ No speech detected ]\n")

                elif msg_type == "cancelled":
                    self._set_transcribing(False)
                    self.status_label.config(text="Transcription cancelled.", fg=WARN_COLOR)

                elif msg_type == "error":
                    self._set_transcribing(False)
                    self.status_label.config(text=f"Error: {data[:80]}", fg=ERROR_COLOR)
                    messagebox.showerror("Transcription Error", str(data))

        except queue.Empty:
            pass

        self.root.after(100, self._poll_queue)

    # ── Output Text Helpers ──────────────────────────────────────────────

    def _set_output(self, text):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete("1.0", tk.END)
        if text:
            self.output_text.insert(tk.END, text)
        self.output_text.config(state=tk.DISABLED)

    def _append_output(self, text):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.insert(tk.END, text)
        self.output_text.see(tk.END)
        self.output_text.config(state=tk.DISABLED)

    # ── Export Actions ───────────────────────────────────────────────────

    def _get_default_stem(self):
        if self.audio_path:
            return sanitize_filename(Path(self.audio_path).stem)
        return "transcription"

    def _export_copy(self):
        if not self.segments:
            return
        text = "\n".join(seg["text"].strip() for seg in self.segments)
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.status_label.config(text="Copied to clipboard!", fg=FG_COLOR)

    def _export_txt(self):
        if not self.segments:
            return
        filepath = filedialog.asksaveasfilename(
            title="Export TXT",
            defaultextension=".txt",
            initialfile=f"{self._get_default_stem()}.txt",
            initialdir=self.prefs.get("last_export_dir", str(Path.home())),
            filetypes=[("Text Files", "*.txt")],
        )
        if not filepath:
            return
        self.prefs["last_export_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)
        export_txt(self.segments, filepath)
        self._verify_export(filepath, "TXT")

    def _export_srt(self):
        if not self.segments:
            return
        filepath = filedialog.asksaveasfilename(
            title="Export SRT",
            defaultextension=".srt",
            initialfile=f"{self._get_default_stem()}.srt",
            initialdir=self.prefs.get("last_export_dir", str(Path.home())),
            filetypes=[("SRT Files", "*.srt")],
        )
        if not filepath:
            return
        self.prefs["last_export_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)
        export_srt(self.segments, filepath)
        self._verify_export(filepath, "SRT")

    def _export_vtt(self):
        if not self.segments:
            return
        filepath = filedialog.asksaveasfilename(
            title="Export VTT",
            defaultextension=".vtt",
            initialfile=f"{self._get_default_stem()}.vtt",
            initialdir=self.prefs.get("last_export_dir", str(Path.home())),
            filetypes=[("WebVTT Files", "*.vtt")],
        )
        if not filepath:
            return
        self.prefs["last_export_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)
        export_vtt(self.segments, filepath)
        self._verify_export(filepath, "VTT")

    def _export_json(self):
        if not self.segments:
            return
        filepath = filedialog.asksaveasfilename(
            title="Export JSON",
            defaultextension=".json",
            initialfile=f"{self._get_default_stem()}.json",
            initialdir=self.prefs.get("last_export_dir", str(Path.home())),
            filetypes=[("JSON Files", "*.json")],
        )
        if not filepath:
            return
        self.prefs["last_export_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)
        export_json(self.segments, filepath, self.metadata)
        self._verify_export(filepath, "JSON")

    def _export_docx(self):
        if not self.segments:
            return
        filepath = filedialog.asksaveasfilename(
            title="Export DOCX",
            defaultextension=".docx",
            initialfile=f"{self._get_default_stem()}.docx",
            initialdir=self.prefs.get("last_export_dir", str(Path.home())),
            filetypes=[("Word Documents", "*.docx")],
        )
        if not filepath:
            return
        self.prefs["last_export_dir"] = str(Path(filepath).parent)
        save_preferences(self.prefs)
        try:
            export_docx(self.segments, filepath, self.metadata)
            self._verify_export(filepath, "DOCX")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export DOCX:\n{e}")

    def _verify_export(self, filepath, fmt):
        """Verify export file exists and is non-zero."""
        if os.path.isfile(filepath) and os.path.getsize(filepath) > 0:
            size = os.path.getsize(filepath)
            self.status_label.config(
                text=f"Exported {fmt}: {os.path.basename(filepath)} ({size:,} bytes)",
                fg=FG_COLOR,
            )
        else:
            self.status_label.config(
                text=f"Export failed — file not created or empty.", fg=ERROR_COLOR,
            )

    # ── Cleanup ──────────────────────────────────────────────────────────

    def _cleanup_temp_files(self):
        for f in self.temp_files:
            try:
                gc.collect()
                if os.path.isfile(f):
                    os.unlink(f)
            except Exception:
                pass

    def _on_close(self):
        if self.worker_thread and self.worker_thread.is_alive():
            self.cancel_event.set()
            self.worker_thread.join(timeout=3)
        self._cleanup_temp_files()
        self.root.destroy()


# ── Entry Point ──────────────────────────────────────────────────────────────

def main():
    root = tk.Tk()
    root.geometry("900x750")

    # Try to set icon (non-critical)
    try:
        root.iconbitmap(default="")
    except Exception:
        pass

    app = AudioScribeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
